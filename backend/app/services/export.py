import io
from docx import Document
from docx.shared import Pt
import fitz


def export_to_docx(resume_text: str) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for line in resume_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        is_header = (
            stripped.isupper()
            or stripped.endswith(":")
            or (len(stripped) < 40 and not stripped.endswith("."))
        )

        if is_header and len(stripped) > 2:
            heading = doc.add_heading(stripped, level=2)
            for run in heading.runs:
                run.font.size = Pt(13)
                run.font.name = "Calibri"
        else:
            if stripped.startswith("-") or stripped.startswith("•"):
                stripped = stripped.lstrip("-• ").strip()
                doc.add_paragraph(stripped, style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_to_pdf(resume_text: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page(width=612, height=792)

    y = 50
    margin_bottom = 750

    for line in resume_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            y += 8
            continue

        is_header = stripped.isupper() or (
            len(stripped) < 40 and not stripped.endswith(".") and not stripped.endswith(",")
        )

        if y > margin_bottom:
            page = doc.new_page(width=612, height=792)
            y = 50

        if is_header and len(stripped) > 2:
            y += 6
            page.insert_text(
                (50, y),
                stripped,
                fontsize=13,
                fontname="helv",
            )
            y += 18
        else:
            wrapped = wrap_text(stripped, max_chars=85)
            for wline in wrapped:
                if y > margin_bottom:
                    page = doc.new_page(width=612, height=792)
                    y = 50
                page.insert_text((50, y), wline, fontsize=10, fontname="helv")
                y += 14
            y += 4

    buffer = io.BytesIO()
    doc.save(buffer)
    doc.close()
    buffer.seek(0)
    return buffer.read()


def wrap_text(text: str, max_chars: int = 85) -> list:
    if len(text) <= max_chars:
        return [text]

    lines = []
    while text:
        if len(text) <= max_chars:
            lines.append(text)
            break
        break_at = text.rfind(" ", 0, max_chars)
        if break_at == -1:
            break_at = max_chars
        lines.append(text[:break_at])
        text = text[break_at:].lstrip()
    return lines
